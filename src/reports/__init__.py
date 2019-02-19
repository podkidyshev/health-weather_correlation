import os
from io import BytesIO

from docx import Document

from science import ScienceError

PATH_DEV = r"src/docx/templates/default.docx"
PATH_EXE = r"default.docx"


def report_error(dest):
    def holy_crap(func):
        def wrapped(self, *args):
            try:
                func(self, *args)
            except Exception as e:
                nonlocal dest
                if dest == "init":
                    dest = "при вычислении данных"
                elif dest == "ui":
                    dest = "при генерации вывода данных на экран"
                elif dest == "doc":
                    dest = "при генерации docx-отчета"
                else:
                    dest = "при генерации отчета"
                raise ScienceError("Непредвиденная ошибка {}\nОтчет: {}\nОшибка: {}"
                                   .format(dest, type(self).__name__, e.args[0]))
        return wrapped
    return holy_crap


def str_arr(arr):
    return "[" + ", ".join(map(lambda x: "{:.2f}".format(x), arr)) + "]"


# noinspection PyUnresolvedReferences,PyTypeChecker
class Printer:
    def __init__(self, destination, func, *args):
        self.destination = destination
        if self.destination == 'doc':
            self.doc = Printer.create_docx()
        elif self.destination == 'ui':
            self.doc = ""
        else:
            raise ValueError("Неизвестное назначение")
        func(*args, self)

    def print(self, destination_obj=None):
        if self.destination == 'ui':
            return self.doc
        else:
            Printer.save_docx(self.doc, destination_obj)
            return True

    def add_heading(self, s, size):
        if self.destination == 'doc':
            self.doc.add_heading(s, size)
        else:
            self.doc += '-- {} --\n\n'.format(s)

    def add_paragraph(self, s):
        if self.destination == 'doc':
            if s[-1] == "\n":
                s = s[:-1]
            self.doc.add_paragraph(s)
        else:
            self.doc += s + '\n'

    def add_picture(self, pic: bytes or bytearray, **kwargs):
        if self.destination == 'doc':
            self.doc.add_picture(BytesIO(pic), **kwargs)

    @staticmethod
    def create_docx():
        doc = Document(PATH_EXE if os.path.exists(PATH_EXE) else PATH_DEV if os.path.exists(PATH_DEV) else None)
        doc._body.clear_content()
        doc.core_properties.author = "Молчанов В.А."
        return doc

    @staticmethod
    def save_docx(doc, obj):
        try:
            doc.save(obj)
        except Exception as e:
            raise ScienceError("Ошибка при сохранении docx-отчета: {}"
                               "\nВозможно уже существующий файл открыт в другой программе".format(e.args[0]))


def print_report(destination, func, *args):
    return Printer(destination, func, *args).print()
