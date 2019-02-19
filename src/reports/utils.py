from docx.shared import Cm

from science import plot_image, FACTORS
from science.funcs import test_normal_plot
from science.classes import Standard, Sample

from reports import Printer, str_arr


def report_ntest(report, doc: Printer):
    res_ok = "пройден"
    res_nok = "не пройден"

    shapiro = report["shapiro"]
    doc.add_paragraph("Тест нормальности Шапиро-Вилка: {}".format(res_ok if shapiro["res"] else res_nok))

    agostino = report["agostino"]
    doc.add_paragraph("Тест нормальности Д'Агостино и Пирсона: из {0} прогонов доля {1}/{0} = {2:.2f} отклоняет "
                      "гипотезу о нормальности на уровне отклонения {3}"
                      .format(agostino["num_tests"], agostino["num_rejects"], agostino["ratio"], agostino["alpha"]))

    ks = report["ks"]
    doc.add_paragraph("Тест нормальности Колмогорова-Смирнова: из {0} прогонов доля {1}/{0} = {2:.2f} отклоняет "
                      "гипотезу о нормальности на уровне отклонения {3}"
                      .format(ks["num_tests"], ks["num_rejects"], ks["ratio"], ks["alpha"]))

    if report['qq'] and doc.destination == "doc":
        img = plot_image(test_normal_plot, report)
        doc.add_paragraph("QQ-тест:")
        doc.add_picture(img, width=Cm(12.5))


def report_stats(stats, doc: Printer):
    doc.add_paragraph("\tВыборочное среднее = {:.2f}".format(stats[0]))
    doc.add_paragraph("\tСтандартное отклонение = {:.2f}".format(stats[1]))
    doc.add_paragraph("\tДоверительный интервал = ({:.2f}, {:.2f})\n".format(*stats[2]))


def report_sample_factor(sample: Sample, factor: int, doc: Printer):
    doc.add_heading("Фактор-образец {}".format(FACTORS[factor].lower()), 2)
    doc.add_paragraph("Количество значений равно = {}".format(len(sample.data[factor])))
    doc.add_paragraph(str_arr(sample.data[factor]))
    doc.add_paragraph("Количество максимумов равно = {}".format(len(sample.seq_max[factor])))
    doc.add_paragraph(str_arr(sample.seq_max[factor]))


def report_sample(sample: Sample, doc: Printer):
    doc.add_heading("Образец {}".format(sample.name), 1)
    for factor in range(4):
        report_sample_factor(sample, factor, doc)


def report_std(std: Standard, doc: Printer):
    doc.add_heading("Эталон {}".format(std.name), 1)
    doc.add_paragraph("Количество значений равно = {}".format(len(std.data)))
    doc.add_paragraph(str_arr(std.data))
    doc.add_paragraph("Количество максимумов равно = {}".format(len(std.seq_max)))
    doc.add_paragraph(str_arr(std.seq_max))
