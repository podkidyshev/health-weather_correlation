import numpy as np
import scipy.stats as stats
import matplotlib.pyplot as plt

from docx import Document

import science
import science.funcs as funcs


def test_normal(x: list, *, qq: bool):
    """Тестирование распределения на нормальность"""
    report = {"x": x[:]}
    alpha = 0.05

    stat, p = stats.shapiro(x)
    report["shapiro"] = {
        "name": "Shapiro-Wilk Test",
        "alpha": alpha,
        "stat": stat,
        "p": p,
        "res": p > alpha
    }

    stat, p = stats.normaltest(x)
    report["agostino"] = {
        "name": "D'Agostino and Pearson's Test",
        "alpha": alpha,
        "stat": stat,
        "p": p,
        "res": p > alpha
    }

    statistic, critical_values, significance_level = stats.anderson(x)
    report["anderson"] = {
        "name": "Anderson-Darling Test",
        "statistic": statistic,
        "critical": critical_values,
        "sig_level": significance_level,
        "res": [statistic < cv for cv in critical_values]
    }

    num_tests = 10 ** (3 if qq else 2)
    num_rejects = 0
    for i in range(num_tests):
        normed_data = (x - np.mean(x)) / np.std(x)
        d, pval = stats.kstest(normed_data, 'norm')
        if pval < alpha:
            num_rejects += 1
    ratio = float(num_rejects) / num_tests
    report["ks"] = {
        "name": "Kolmogorov-Smirnov Test",
        "num_tests": num_tests,
        "num_rejects": num_rejects,
        "ratio": ratio,
        "alpha": alpha
    }
    return report


def get_report(report, doc: Document):
    res_ok = "Образец выглядит гауссовским (не может отклонить гипотезу H0)"
    res_nok = "Образец не выглядит гауссовским (отклонить гипотезу H0)"

    shapiro = report["shapiro"]
    doc.add_heading("Тест нормальности Шапиро-Вилка", 2)
    doc.add_paragraph("Statistics = {:.3f}, p = {:.3f}".format(shapiro["stat"], shapiro["p"]))
    doc.add_paragraph(res_ok if shapiro["res"] else res_nok)

    agostino = report["agostino"]
    doc.add_heading("D'Agostino and Pearson's Test", 2)
    doc.add_paragraph("Statistics = {:.3f}, p = {:.3f}".format(agostino["stat"], agostino["p"]))
    doc.add_paragraph(res_ok if agostino["res"] else res_nok)

    anderson = report["anderson"]
    doc.add_heading("Тест нормальности Андерсона-Дарлинга", 2)
    doc.add_paragraph("Statistic = {:.3f}".format(anderson["statistic"]))
    for res, cv, sl in zip(anderson["res"], anderson["critical"], anderson["sig_level"]):
        doc.add_paragraph("{:.3f}: {:.3f}, {}".format(sl, cv, res_ok if res else res_nok))

    ks = report["ks"]
    doc.add_heading("Тест нормальности Колмогорова-Смирнова", 2)
    num_tests = ks["num_tests"]
    num_rejects = ks["num_rejects"]
    ratio = ks["ratio"]
    alpha = ks["alpha"]
    doc.add_paragraph(
        "Результаты теста Колмогорова-Смирнова: "
        "из {} прогонов доля {}/{} = {:.2f} отклоняет гипотезу H0 на уровне отклонения {}".format(
            num_tests, num_rejects, num_tests, ratio, alpha))
    base = plt.figure()
    get_plot(report, base)
    doc.add_picture(science.plot_to_stream(base))


def get_plot(report, base):
    fig = base.subplots(1, 1)
    stats.probplot(report["x"], dist="norm", plot=fig)


def test():
    x_distance = [1, -1, 1, 1, -3, 2, 0, 2, -2, -1, 0, 1, 1, 3, -3, 0, 1, -1, 3]

    report = test_normal(x_distance, qq=True)

    doc = science.create_docx()
    get_report(report, doc)
    science.save_docx(doc, "../test_normal.docx")


if __name__ == '__main__':
    test()
