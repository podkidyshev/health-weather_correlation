# -*- coding: utf-8 -*-
# Ввод образцов_послед.максимумов_распред..расстояний_гистограммы
import numpy as np
import scipy.stats as stats

import matplotlib.pyplot as plt

from docx import Document

from science.classes import Patient, Standard
from science.funcs import sequence_distance, distrib, graph_kde
from science import CATS, nnone, plot_to_stream


def std_pat_stat_by_category(std: Standard, pat: Patient, cat: int):
    distance = sequence_distance(pat.seq_max[cat], std.seq_max, insert_zero=True)
    return {
        "seq_max": pat.seq_max[cat],
        "distance": distance,
        "distrib": distrib(distance),
        'mean': np.mean(distance),
        'std': np.std(distance),
        't-interval': stats.t.interval(0.95, len(distance) - 1, loc=np.mean(distance), scale=stats.sem(distance))
    }


def st_pat_stat(std: Standard, pat: Patient):
    report = [None] * len(CATS)
    for cat, data in nnone(pat.data):
        report[cat] = std_pat_stat_by_category(std, pat, cat)
    return report


class StandardPatientStat:
    def __init__(self, std: Standard, pat: Patient):
        self.std = std
        self.pat = pat

        self.report = st_pat_stat(std, pat)

    def get_report_item(self, item: str):
        return [self.report[idx][item] if self.report[idx] is not None else None for idx in range(len(CATS))]

    def get_report(self, doc: Document):
        doc.add_heading("Пациент {}. Эталон {}".format(self.pat.name, self.std.name), 0)

        doc.add_heading("Список значений эталона", 1)
        doc.add_paragraph("Количество значений равно = {}".format(len(self.std.data)))
        doc.add_paragraph(str(self.std.data))

        doc.add_heading("Список максимумов Кр-значений:", 1)
        doc.add_paragraph("Количество значений равно = {}".format(str(len(self.std.seq_max))))
        doc.add_paragraph(str(self.std.seq_max))

        doc.add_paragraph()
        for cat, report in nnone(self.report):
            doc.add_paragraph("Список значений пациента {}:".format(CATS[cat][1]))
            doc.add_paragraph("Количество значений равно = {}".format(len(self.pat.data[cat])))
            doc.add_paragraph(str(self.pat.data[cat]))

            doc.add_paragraph("Список максимумов значений пациента {}:".format(CATS[cat][1]))
            doc.add_paragraph("Количество значений равно = {}".format(len(self.pat.seq_max[cat])))
            doc.add_paragraph(str(self.pat.seq_max[cat]))
            doc.add_paragraph()

        # вычисление распределения расстояний от максимумов рядов пациентов до ближайшего максимума эталона
        # вычисление последовательностей расстояний  от максимумов рядов пациентов до ближайшего максимума эталона
        doc.add_paragraph()
        doc.add_heading(
            "Ряды расстояний и распределения расстояний от максимумов пациента до ближайшего максимума эталона", 1)
        for cat, report in nnone(self.report):
            doc.add_paragraph(
                "Ряд расстояний от максимумов пациента {} до ближайшего максимума эталона:".format(CATS[cat][1]))
            doc.add_paragraph(str(report["distance"]))
            doc.add_paragraph("Распределение расстояний (значения от -3 до 3) пациента {}".format(CATS[cat][1]))
            doc.add_paragraph(str(report["distrib"]))
            doc.add_paragraph()

        doc.add_heading("Анализ распределений расстояний от максимумов пациента до ближайшего максимума эталона", 1)
        for cat, report in nnone(self.report):
            doc.add_paragraph("Анализ распределений расстояний пациента {}:".format(CATS[cat][1]))
            doc.add_paragraph("\tвыборочное среднее = {:.4f}".format(report["mean"]))
            doc.add_paragraph("\tстандартное отклонение = {:.4f}".format(report["std"]))
            doc.add_paragraph("\tдоверительный интервал = ({:.4f}, {:.4f})".format(*report["t-interval"]))

        doc.add_heading('График анализа:', 1)
        base = plt.figure()
        graph_kde(self.get_report_item("distance"), base)
        doc.add_picture(plot_to_stream(base))


def test():
    from science import create_docx, save_docx, plot_to_image

    std = Standard.from_file('samples\\Flow_62.txt')

    pat = Patient.from_file("samples/1_1.xlsx")
    stat = StandardPatientStat(std, pat)

    doc = create_docx()
    stat.get_report(doc)
    save_docx(doc, "test.docx")

    plt.clf()

    base = plt.figure()
    graph_kde(stat.get_report_item("distance"), base)
    plt.show()


if __name__ == '__main__':
    test()
