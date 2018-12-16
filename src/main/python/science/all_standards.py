import numpy as np
import scipy.stats as stats
import matplotlib.pyplot as plt

from docx import Document

import science
import science.test_normal as tn
import science.funcs as funcs
import science.classes as classes


class AllStandards:
    def __init__(self, samples: list, cat: int):
        """Анализ по всем эталонам и переданным пациентам в категории cat"""
        self.cat = cat
        self.cat_name = science.CATS[cat]

        error_samples = []
        for sample in samples:
            if not sample.has_cat(cat):
                error_samples.append(sample.name)
        if len(error_samples):
            raise funcs.StatComputingError('Запрошена обработка категории "{}".\n '
                                           'У пациентов {} нет данных в этой категории'.format(self.cat_name,
                                                                                               error_samples))

        self.stds = [std for std in classes.Standard.standards.values()]
        self.stds_count = len(self.stds)
        self.stds_data = [std.data for std in self.stds]
        self.stds_seq_max = [std.seq_max for std in self.stds]

        self.pats = samples[:]
        self.pats_count = len(self.pats)
        self.pats_data = [pat.data[self.cat] for pat in self.pats]
        self.pats_seq_max = [pat.seq_max[self.cat] for pat in self.pats]

        self.pats_data_all = funcs.sum_list(self.pats_data)
        self.pats_seq_max_all = funcs.sequence_max(self.pats_data_all)

        self.distances = [[funcs.sequence_distance(pat_seq_max, std_seq_max, insert_zero=True)
                           for pat_seq_max in self.pats_seq_max]
                          for std_seq_max in self.stds_seq_max]

        self.distances_all = [funcs.sequence_distance(self.pats_seq_max_all, std_seq_max, insert_zero=True)
                              for std_seq_max in self.stds_seq_max]

        self.concats = [funcs.sum_list(self.distances[i]) for i in range(self.stds_count)]
        self.tn_by_day = [tn.test_normal(funcs.sequence_distance(self.pats_seq_max_all,
                                                                 std_seq_max,
                                                                 insert_zero=True),
                                         qq=False)
                          for std_seq_max in self.stds_seq_max]
        self.tn_by_pat = [tn.test_normal(concat, qq=False) for concat in self.concats]

    def get_report(self, doc: Document):
        doc.add_heading("Анализ по всем эталонам по всем пациентам {}".format(self.cat_name), 0)

        doc.add_paragraph("Число эталонов = {}\n".format(self.stds_count))
        for std in self.stds:
            doc.add_paragraph("Эталон {}".format(std.name))
            doc.add_paragraph("Список значений: {}".format(std.data))
            doc.add_paragraph("Всего максимумов = {}".format(np.sum(std.seq_max)))
            doc.add_paragraph("Список максимумов значений: {}\n".format(std.seq_max))

        doc.add_paragraph("Число образцов = {}".format(self.pats_count))

        doc.add_heading("Распределения максимумов и расстояний по пациентам для всех эталонов", 1)
        for sidx, std in enumerate(self.stds):
            for pidx, pat in enumerate(self.pats):
                doc.add_paragraph(
                    "Последовательность расстояний для образца {} и эталона {}. Количество элементов = {}".format(
                        pat.name, std.name, len(self.distances[sidx][pidx])))

        doc.add_heading("Распределения расстояний образцов по эталонам", 1)
        for std, concat in zip(self.stds, self.concats):
            doc.add_paragraph("Распределение расстояний образцов для эталона {}".format(std.name))
            doc.add_paragraph("Количество значений равно = {}".format(len(concat)))
            doc.add_paragraph("Значения: {}\n".format(concat))

        doc.add_heading(
            "Результаты сравнительного визуального анализа по дням и по пациентам всех образцов со всеми эталонами", 1)
        for std_idx in range(self.stds_count):
            doc.add_heading("Эталон {}".format(self.stds[std_idx].name), 2)
            base = plt.figure()
            funcs.visual_analysis2(funcs.sequence_distance(funcs.sequence_max(self.pats_data_all),
                                                           self.stds_seq_max[std_idx], insert_zero=True),
                                   self.concats[std_idx], base)
            doc.add_picture(science.plot_to_stream(base))

            doc.add_paragraph("Результаты тестирования нормальности распределения группового образца по дням")
            tn.get_report(self.tn_by_day[std_idx], doc)
            doc.add_paragraph("Результаты тестирования нормальности распределения группового образца по пациентам")
            tn.get_report(self.tn_by_pat[std_idx], doc)

        doc.add_heading("Результаты группового анализа по дням", 1)
        doc.add_paragraph("[Выборочное среднее, Стандартное отклонение,  Доверительный интервал] =")
        for std_idx in range(self.stds_count):
            doc.add_heading("Для эталона {}".format(self.stds[std_idx].name), 2)
            doc.add_paragraph("Результаты группового анализа по дням")
            doc.add_paragraph(str(funcs.stat_analysis_distances(self.pats_data_all, self.stds_data[std_idx])))
            doc.add_paragraph("Результаты группового анализа по пациентам")
            doc.add_paragraph(str(funcs.stat_analysis(self.concats[std_idx])))

        doc.add_heading("Визуализация данных для группового образца по дням: гистограмма, "
                        "ядерная оценка плотности распределения расстояний, плотность нормального распределения", 1)
        for idx, std in enumerate(self.stds):
            doc.add_heading("Для эталона {}".format(std.name), 2)
            base = plt.figure()
            funcs.visual_analysis(self.distances_all[idx], base)
            doc.add_picture(science.plot_to_stream(base))

        doc.add_heading("Результаты статистического анализа", 1)
        for sidx, std in enumerate(self.stds):
            doc.add_heading("Для эталона {}".format(std.name), 2)
            for pidx, pat in enumerate(self.pats):
                data = self.distances[sidx][pidx]
                doc.add_paragraph("Выборочное среднее = {:.3f}".format(np.mean(data)))
                doc.add_paragraph("Стандартное отклонение = {:.3f}".format(np.std(data)))
                doc.add_paragraph("Доверительный интервал = {}".format(stats.t.interval(0.95, len(data)-1,
                                                                                        loc=np.mean(data),
                                                                                        scale=stats.sem(data))))

                base = plt.figure()
                funcs.visual_analysis(data, base)
                doc.add_picture(science.plot_to_stream(base))


def test():
    from science.classes import Patient, Standard
    Standard.from_file("samples/Flow_62.txt")
    Standard.from_file("samples/Kp_62.txt")

    pat1 = Patient.from_file("samples/1_1.xlsx")
    pat2 = Patient.from_file("samples/1_2.xlsx")

    stat = AllStandards([pat1, pat2], 0)

    doc = science.create_docx()
    stat.get_report(doc)
    science.save_docx(doc, "../test_all.docx")


if __name__ == '__main__':
    test()
