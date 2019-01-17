from io import BytesIO

from docx import Document
from docx.shared import Cm

import science
import science.funcs
import science.classes


def str_arr(arr):
    return "[" + ", ".join(map(lambda x: "{:.2f}".format(x), arr)) + "]"


# noinspection PyUnresolvedReferences,PyTypeChecker
class Printer:
    def __init__(self, destination, func, *args):
        self.destination = destination
        if self.destination == 'doc':
            self.doc = Printer.create_docx()
        elif self.destination == 'ui':
            self.doc = ""
        else:
            raise ValueError("Неизвестное назначение")
        func(*args, self)

    def print(self, destination_obj=None):
        if self.destination == 'ui':
            return self.doc
        else:
            Printer.save_docx(self.doc, destination_obj)
            return True

    def add_heading(self, s, size):
        if self.destination == 'doc':
            self.doc.add_heading(s, size)
        else:
            self.doc += '-- {} --\n\n'.format(s)

    def add_paragraph(self, s):
        if self.destination == 'doc':
            self.doc.add_paragraph(s)
        else:
            self.doc += s + '\n'

    def add_picture(self, pic: bytes or bytearray):
        if self.destination == 'doc':
            self.doc.add_picture(BytesIO(pic))

    @staticmethod
    def create_docx():
        doc = Document()
        doc.core_properties.author = "Молчанов В.А."
        return doc

    @staticmethod
    def save_docx(doc, obj):
        # TODO: убрать, настроить default-паттерн в src/docx
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)
        doc.save(obj)


def print_report(destination, func, *args):
    return Printer(destination, func, *args).print()
